"""
Generate the 10-document Incident Operations knowledge-base corpus.

Run once:
    python scripts/build_corpus.py

Output: projects/piter-aiops/data/sample_documents/
  Formats covered: MD, TXT, CSV, DOCX, PDF (per assignment spec).

These files are checked into the repo so the Bedrock KB can be rebuilt
from scratch via `infra/upload_docs_to_s3.sh` and a KB sync.
"""
from __future__ import annotations

import csv
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors

OUT_DIR = Path(__file__).resolve().parents[1] / "data" / "sample_documents"
OUT_DIR.mkdir(parents=True, exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# 1) auth_service_runbook.md
# ─────────────────────────────────────────────────────────────────────────────
def write_auth_runbook_md() -> None:
    content = dedent("""\
        # Authentication Service Incident Runbook

        **Owner:** Identity & Access Team
        **Severity classes covered:** P1 (total outage), P2 (partial), P3 (degraded)
        **Last reviewed:** 2026-02-14

        ## When to use this runbook

        Use this runbook for any incident where users cannot authenticate, OIDC/SAML
        flows are failing, MFA challenges are stuck, or session tokens are being
        rejected by downstream services.

        ## Post-deployment login failures (FAQ)

        **Question:** What should I check first when users cannot log in after a deployment?

        **Answer:** Follow the first-response checklist below. When login breaks right
        after a rollout, **check recent deployments first** (step 2) before debugging
        Redis, Postgres, or OIDC. If the latest `auth-api` rollout is less than 30
        minutes old, treat the deployment as the likely cause and prepare a rollback
        per the Standard recovery actions section.

        ## First-response checklist (5 minutes)

        1. **Confirm scope.** Check the auth service dashboard for the global error
           rate. If `auth_login_errors_total` is above 5% over the last 5 minutes,
           declare a P2 immediately. Above 25% → P1.
        2. **Check recent deployments.** `kubectl -n auth get deploy -o wide` —
           if the latest rollout is < 30 minutes old, treat it as the likely cause.
        3. **Check upstream dependencies.** The auth service depends on Redis
           (session store) and Postgres (user table). Confirm both are healthy in
           the Service Map view before suspecting auth itself.
        4. **Validate the OIDC discovery document.** `curl -sS https://auth.example.com/.well-known/openid-configuration | jq`.
           A 5xx or empty response means the public ingress is broken — escalate
           to the platform team.

        ## Triage decision tree

        - **Symptom: login page returns 502/504**
          → Likely ingress or pod readiness. Run `kubectl -n auth get pods` and
          check the `auth-api` deployment for crashlooping replicas.
        - **Symptom: login returns 401 for valid credentials**
          → Likely session store. Run the Redis health command in the next
          section. If Redis is degraded, follow the database connectivity runbook.
        - **Symptom: MFA challenges never arrive**
          → Likely SMS provider or email provider. Check the provider status
          page first; failover provider is documented in `escalation_policy.pdf`.

        ## Standard recovery actions

        1. **Roll back the latest deployment** if it was deployed in the last hour:
           `kubectl -n auth rollout undo deployment/auth-api`.
        2. **Restart the pod fleet** if a rollback is not appropriate:
           `kubectl -n auth rollout restart deployment/auth-api`.
        3. **Clear corrupted session entries** if users report being logged out
           mid-session: see `database_connectivity_runbook.md` for the Redis FLUSHDB
           procedure (only with on-call lead approval).
        4. **Verify recovery** by hitting `/health/ready` on each pod and watching
           the `auth_login_success_total` counter return to baseline (~99.5%).

        ## Communication

        Post in `#incident-auth` every 15 minutes with: current error rate,
        actions taken, ETA, customer impact estimate.
    """)
    (OUT_DIR / "auth_service_runbook.md").write_text(content, encoding="utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# 2) database_connectivity_runbook.md
# ─────────────────────────────────────────────────────────────────────────────
def write_db_runbook_md() -> None:
    content = dedent("""\
        # Database Connectivity Runbook

        Applies to: PostgreSQL primary + replicas, Redis session/cache layer.

        ## Symptoms that map here

        - Applications log `connection refused`, `too many connections`, or
          `SSL handshake failed`
        - p95 query latency above 1 second sustained for 5 minutes
        - Replication lag above 30 seconds on any read replica
        - `idle in transaction` connections above 50

        ## Postgres — fast checks

        ```bash
        # From a bastion host
        psql -h prod-pg-primary.example.com -U readonly -c \\
          "SELECT count(*) FROM pg_stat_activity WHERE state='active';"
        psql ... -c "SELECT now() - pg_last_xact_replay_timestamp() AS lag;"
        ```

        If active connections are above 80% of `max_connections`, the application
        is leaking. Restart the offending app pod fleet **before** touching the
        database. Killing Postgres backends only buys a few seconds.

        ## Redis — fast checks

        ```bash
        redis-cli -h prod-redis.example.com PING
        redis-cli -h prod-redis.example.com INFO replication | head
        redis-cli -h prod-redis.example.com INFO memory | grep used_memory_human
        ```

        If `used_memory_human` is within 90% of the maxmemory cap, evictions
        will start failing session writes — see auth runbook.

        ## When the database is the actual cause

        - **Locks:** `SELECT * FROM pg_locks WHERE NOT granted;` — kill the
          oldest blocking PID only with DBA on the call.
        - **Replica lag:** if a read replica is more than 5 minutes behind, take
          it out of rotation in the load balancer.
        - **Disk full:** `df -h /var/lib/postgresql` — if above 90%, immediately
          truncate the `audit_log` table after exporting; this needs the DBA.

        ## Do NOT

        - Do **not** run `FLUSHDB` against Redis without on-call lead approval.
        - Do **not** restart Postgres during business hours without a documented
          recovery window — failover takes ~3 minutes and drops in-flight
          transactions.
    """)
    (OUT_DIR / "database_connectivity_runbook.md").write_text(content, encoding="utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# 3) monitoring_alerts_reference.md
# ─────────────────────────────────────────────────────────────────────────────
def write_monitoring_md() -> None:
    content = dedent("""\
        # Monitoring & Alert Reference

        A single source of truth for what each alert means, what the on-call
        engineer should do first, and which runbook to consult.

        ## P1 — page immediately

        | Alert name | Condition | First action | Runbook |
        |---|---|---|---|
        | `AuthGlobalErrorRateCritical` | login error rate > 25% for 5m | Roll back if recent deploy; else escalate | `auth_service_runbook.md` |
        | `PaymentGatewayDown` | 100% non-2xx for 2m on `/charge` | Failover to backup PSP | `payment_service_latency_runbook.txt` |
        | `DatabasePrimaryUnreachable` | health probe fails 3× | Promote standby; involve DBA | `database_connectivity_runbook.md` |
        | `ApiGateway5xxStorm` | 5xx > 10% for 3m | Check upstream service map | `api_gateway_5xx_runbook.txt` |

        ## P2 — page within 15 minutes

        | Alert | Condition | First action |
        |---|---|---|
        | `AuthMfaProviderDegraded` | MFA delivery success < 95% for 10m | Failover to backup provider |
        | `ReplicationLagHigh` | Postgres lag > 60s for 10m | Remove replica from LB |
        | `DiskUsageHigh` | any volume > 85% for 15m | Rotate logs / extend volume |

        ## P3 — handle during business hours

        Anything labelled `severity:warn` in Prometheus. These do **not** page
        out of hours unless they roll up into a P1/P2.

        ## Silence policy

        Silences must always include a Jira ticket reference and an expiry of
        at most 24 hours. Open-ended silences are forbidden.
    """)
    (OUT_DIR / "monitoring_alerts_reference.md").write_text(content, encoding="utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# 4) api_gateway_5xx_runbook.txt
# ─────────────────────────────────────────────────────────────────────────────
def write_apigw_txt() -> None:
    content = dedent("""\
        API GATEWAY — 5xx STORM RUNBOOK
        ===============================

        Trigger: AWS API Gateway returns 5xx for more than 10% of requests over
        a 3-minute window, OR the CloudWatch alarm "ApiGateway5xxStorm" fires.

        STEP 1 — Identify the failing route
        -----------------------------------
        Open CloudWatch -> API Gateway -> select the stage. Look at the "5XX by
        resource" chart for the last 15 minutes. The top 1-2 routes account for
        most failures in 90% of incidents.

        STEP 2 — Decide: gateway issue or backend issue?
        ------------------------------------------------
        Compare the gateway 5xx rate with the integration target's own 5xx rate
        (Lambda errors, ECS task errors, or upstream HTTP service).
          - If backend errors mirror gateway errors -> the backend is the cause,
            jump to that service's runbook.
          - If backend is healthy but gateway still 5xxs -> usage plan throttle
            or integration timeout. Continue to step 3.

        STEP 3 — Check throttling
        -------------------------
        API Gateway -> Usage Plans -> verify that no plan is at its throttle
        limit. If a known customer is exceeding their quota, that is expected;
        document and move on.

        STEP 4 — Check integration timeout
        ----------------------------------
        Default integration timeout is 29 seconds. If backend latency has crept
        above 25 seconds, gateway will return 504. Roll back the backend if a
        recent deploy is suspect, OR scale the backend horizontally.

        STEP 5 — Recovery validation
        ----------------------------
        After the fix, monitor for 10 minutes:
          - 5xx rate must drop below 1%
          - p99 latency must return to baseline
          - No new alerts firing for the same route

        DO NOT
        ------
        Do not delete deployment stages to "reset" the gateway. This severs all
        custom domain mappings and triggers a much larger incident.
    """)
    (OUT_DIR / "api_gateway_5xx_runbook.txt").write_text(content, encoding="utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# 5) payment_service_latency_runbook.txt
# ─────────────────────────────────────────────────────────────────────────────
def write_payment_txt() -> None:
    content = dedent("""\
        PAYMENT SERVICE — LATENCY RUNBOOK
        =================================

        Trigger conditions
        ------------------
        - p95 of POST /charge above 2000 ms for 5 minutes
        - or, success rate below 99.0% for 5 minutes
        - or, an alert from the primary payment provider's status page

        Initial triage (under 3 minutes)
        --------------------------------
        1. Check the provider status page. If the primary PSP is degraded, skip
           directly to the failover step.
        2. Check internal metrics: charge_attempts_total, charge_success_total,
           charge_p95_latency_ms.
        3. If only ONE region is affected, traffic-shift to the healthy region
           via the global load balancer. Document the shift in #incident-pay.

        Failover to the backup PSP
        --------------------------
        We support two PSPs: PRIMARY=stripe, BACKUP=adyen. Failover is a feature
        flag, not a deployment:

            curl -X POST https://configsvc.example.com/flags/payment.psp \\
                 -H 'Authorization: Bearer $CONFIG_TOKEN' \\
                 -d '{"value":"adyen"}'

        Effect is immediate. Verify by watching charge_attempts_total{psp="adyen"}
        increase within 30 seconds. Roll back the flag once the primary recovers
        and has been stable for 30 minutes.

        DO NOT
        ------
        - Do not retry failed charges automatically without idempotency keys.
          We have caused duplicate-charge incidents before. The retry helper in
          payments-sdk handles this safely; never write your own retry loop.

        Communication
        -------------
        Customer-facing status page MUST be updated within 10 minutes of declaring
        a payment incident. Use the "Payments degraded" template.
    """)
    (OUT_DIR / "payment_service_latency_runbook.txt").write_text(content, encoding="utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# 6) incident_history.csv
# ─────────────────────────────────────────────────────────────────────────────
def write_incident_csv() -> None:
    rows = [
        ["incident_id", "date", "severity", "service", "root_cause", "mttr_minutes", "customer_impact"],
        ["INC-2025-001", "2025-01-08", "P1", "auth-api", "Bad config rollout disabled session cookie", 47, "100% of users unable to log in for 47 min"],
        ["INC-2025-002", "2025-01-19", "P2", "payment-svc", "Primary PSP timeout spike", 31, "12% of charges retried via backup PSP"],
        ["INC-2025-003", "2025-02-02", "P3", "api-gateway", "Single Lambda integration timeout", 18, "One customer endpoint slow"],
        ["INC-2025-004", "2025-02-14", "P1", "postgres-primary", "Disk full from runaway audit log", 62, "All writes failing for 62 min"],
        ["INC-2025-005", "2025-02-28", "P2", "auth-api", "MFA SMS provider degraded", 25, "MFA delivery delayed up to 5 min"],
        ["INC-2025-006", "2025-03-10", "P3", "monitoring", "Prometheus OOM killed", 12, "Alerts delayed but not lost"],
        ["INC-2025-007", "2025-03-22", "P1", "payment-svc", "Duplicate charge bug from custom retry loop", 95, "0.4% of customers double-charged"],
        ["INC-2025-008", "2025-04-03", "P2", "redis-cache", "Eviction storm under load", 28, "Sessions expiring early"],
        ["INC-2025-009", "2025-04-15", "P3", "api-gateway", "Throttle limit reached for one client", 9, "One enterprise client rate-limited"],
        ["INC-2025-010", "2025-04-29", "P1", "auth-api", "OIDC discovery endpoint cached stale keys", 73, "All OAuth logins failing"],
        ["INC-2025-011", "2025-05-12", "P2", "postgres-replica", "Replica fell 8 min behind", 22, "Read-after-write inconsistencies"],
        ["INC-2025-012", "2025-05-25", "P3", "deploy-pipeline", "CI flakiness blocked release", 41, "Deploy delayed, no customer impact"],
        ["INC-2025-013", "2025-06-07", "P1", "api-gateway", "Custom domain cert expired", 18, "All public APIs unreachable"],
        ["INC-2025-014", "2025-06-20", "P2", "auth-api", "Token signing key rotation race", 33, "5% of new sessions rejected"],
        ["INC-2025-015", "2025-07-04", "P3", "monitoring", "Grafana datasource auth expired", 7, "Dashboards blank, no real impact"],
        ["INC-2025-016", "2025-07-17", "P1", "payment-svc", "PSP webhook signing key leaked, rotated emergency", 56, "Webhook ingestion stopped for 56 min"],
        ["INC-2025-017", "2025-07-30", "P2", "auth-api", "Pod CrashLoop from OOM after JDK upgrade", 24, "Login error rate 8%"],
        ["INC-2025-018", "2025-08-11", "P3", "api-gateway", "WAF false-positive blocked one customer", 14, "Single customer 403s"],
        ["INC-2025-019", "2025-08-24", "P1", "postgres-primary", "Connection pool exhausted from leaking app", 41, "All writes failing intermittently"],
        ["INC-2025-020", "2025-09-06", "P2", "payment-svc", "Currency rate cache stale", 36, "Wrong currency conversion on 0.2% of orders"],
        ["INC-2025-021", "2025-09-19", "P3", "auth-api", "MFA email provider rate-limited", 11, "Email MFA delayed"],
        ["INC-2025-022", "2025-10-02", "P1", "redis-cache", "Master node OOM, failover incomplete", 39, "Sessions invalidated"],
        ["INC-2025-023", "2025-10-15", "P2", "api-gateway", "Stage variable typo broke routing", 19, "Beta endpoint 5xx"],
        ["INC-2025-024", "2025-10-28", "P3", "monitoring", "Alert silence forgotten 6 hours", 360, "No customer impact, missed minor alerts"],
        ["INC-2025-025", "2025-11-10", "P1", "auth-api", "Bad migration locked users table 4 min", 16, "All writes blocked, reads OK"],
        ["INC-2025-026", "2025-11-23", "P2", "payment-svc", "Webhook backlog after PSP recovery", 47, "Order confirmation delayed up to 30 min"],
        ["INC-2025-027", "2025-12-06", "P3", "deploy-pipeline", "Docker registry quota hit", 28, "No customer impact, deploys blocked"],
        ["INC-2025-028", "2025-12-19", "P1", "api-gateway", "DDoS, WAF rate limits insufficient", 88, "30% of public traffic dropped"],
        ["INC-2025-029", "2026-01-01", "P2", "auth-api", "New year leap-second clock skew", 21, "Some tokens marked invalid"],
        ["INC-2025-030", "2026-01-14", "P3", "monitoring", "Prometheus disk near full", 17, "No customer impact"],
    ]
    with (OUT_DIR / "incident_history.csv").open("w", newline="") as f:
        csv.writer(f).writerows(rows)


# ─────────────────────────────────────────────────────────────────────────────
# 7) deployment_rollback_sop.docx
# ─────────────────────────────────────────────────────────────────────────────
def write_deployment_docx() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Deployment Rollback — Standard Operating Procedure", level=0)
    doc.add_paragraph(
        "This SOP defines how on-call engineers safely roll back a production "
        "deployment that is causing or suspected of causing a live incident. "
        "It applies to all services deployed via the standard CI/CD pipeline."
    )

    doc.add_heading("When to roll back", level=1)
    doc.add_paragraph(
        "Roll back immediately when ALL of the following are true:"
    )
    for item in [
        "A new deployment shipped within the last 60 minutes.",
        "Error rate, latency, or success-rate metrics have degraded since that deployment.",
        "No simpler mitigation (feature-flag, config reversal) is available within 5 minutes.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Do NOT roll back if the deployment is older than 24 hours; the bug is "
        "more likely to be in data, configuration, or upstream dependencies. "
        "Investigate before reverting code in that case."
    )

    doc.add_heading("Rollback procedure (Kubernetes services)", level=1)
    steps = [
        "Open the on-call dashboard and confirm the current deployment SHA.",
        "Run `kubectl -n <namespace> rollout history deployment/<name>` and identify the previous revision.",
        "Announce the rollback in `#incident-<service>` with the from-SHA and to-SHA.",
        "Execute `kubectl -n <namespace> rollout undo deployment/<name>`.",
        "Watch `kubectl -n <namespace> rollout status deployment/<name>` until all pods are Ready.",
        "Verify metrics for 10 minutes before declaring the incident mitigated.",
    ]
    for i, step in enumerate(steps, start=1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("Rollback procedure (Lambda functions)", level=1)
    doc.add_paragraph(
        "Each Lambda alias points to a specific function version. Move the "
        "production alias back to the previous version with the AWS CLI:"
    )
    doc.add_paragraph(
        "aws lambda update-alias --function-name <name> "
        "--name prod --function-version <previous-version>"
    )

    doc.add_heading("After the rollback", level=1)
    for item in [
        "File an incident ticket with the rollback timeline.",
        "Open a follow-up PR in the source repo that includes a regression test.",
        "Add the failure mode to the on-call handoff checklist if it is not already there.",
        "Schedule a postmortem within 5 business days using the standard template.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Anti-patterns we have caused production incidents with", level=1)
    for item in [
        "Rolling back the application but not the corresponding database migration.",
        "Skipping the metric-verification window and declaring resolved too early.",
        "Rolling back during a deployment freeze window without on-call lead approval.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(OUT_DIR / "deployment_rollback_sop.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 8) postmortem_template.docx
# ─────────────────────────────────────────────────────────────────────────────
def write_postmortem_docx() -> None:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("Incident Postmortem — Template", level=0)
    doc.add_paragraph(
        "Every P1 and P2 incident requires a postmortem within 5 business days "
        "of resolution. Postmortems are blameless: focus on systems, processes, "
        "and contributing factors, not on individuals."
    )

    sections = [
        ("Incident summary", "One paragraph: what happened, who was affected, for how long."),
        ("Timeline", "Bullet list with timestamps in UTC. Include detection, first response, mitigation, and resolution."),
        ("Impact", "Customer-facing impact in measurable terms: requests failed, users affected, revenue at risk."),
        ("Root cause", "The actual technical cause. If there are multiple contributing factors, list each."),
        ("Detection", "How was the incident detected? Was the alert timely? Was the right team paged?"),
        ("Response", "What worked well during the response? What slowed us down?"),
        ("Action items", "Bullet list of concrete, owned, dated follow-ups. Each item must have an owner and a due date."),
        ("Lessons learned", "What patterns or assumptions does this incident challenge? Update runbooks accordingly."),
    ]
    for title, body in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)

    doc.add_heading("Action item discipline", level=1)
    for item in [
        "No action item without an owner.",
        "No action item without a due date.",
        "No vague items like 'improve monitoring'; instead write 'add alert for X above Y for Z minutes'.",
        "Track all items in the incidents Jira board; revisit weekly.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(OUT_DIR / "postmortem_template.docx")


# ─────────────────────────────────────────────────────────────────────────────
# 9) escalation_policy.pdf
# ─────────────────────────────────────────────────────────────────────────────
def write_escalation_pdf() -> None:
    out = OUT_DIR / "escalation_policy.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=LETTER,
                            leftMargin=0.75*inch, rightMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=11, leading=15, spaceAfter=8)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=18, spaceAfter=12)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13, spaceAfter=8, spaceBefore=10)

    flow = [
        Paragraph("Incident Escalation Policy", h1),
        Paragraph("Owner: Site Reliability Engineering. Reviewed quarterly. "
                  "This policy applies to all production-facing services.", body),

        Paragraph("Severity definitions", h2),
    ]

    sev_data = [
        ["Severity", "Definition", "Page", "Status page", "Postmortem"],
        ["P1", "Total outage or customer data risk", "Immediately", "Within 15 min", "Required"],
        ["P2", "Significant degradation, partial outage", "Within 15 min", "Within 30 min", "Required"],
        ["P3", "Minor degradation, single feature", "Business hours", "Not required", "Optional"],
        ["P4", "Cosmetic, no customer impact", "Business hours", "Not required", "Not required"],
    ]
    t = Table(sev_data, colWidths=[0.7*inch, 2.6*inch, 1.1*inch, 1.1*inch, 1.1*inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#0b1220")),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("ALIGN", (0,0), (-1,-1), "LEFT"),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("BOTTOMPADDING", (0,0), (-1,-1), 6),
        ("TOPPADDING", (0,0), (-1,-1), 6),
    ]))
    flow.append(t)
    flow.append(Spacer(1, 12))

    flow.append(Paragraph("On-call chain", h2))
    flow.append(Paragraph(
        "Primary on-call carries the page. If no acknowledgement within 8 minutes, "
        "the secondary is paged automatically. If secondary does not ack within "
        "8 more minutes, the engineering manager on call is paged.", body))

    flow.append(Paragraph("When to wake the engineering manager", h2))
    flow.append(Paragraph(
        "Wake the EM on call for: any P1 lasting more than 30 minutes; any "
        "incident where customer data confidentiality, integrity, or availability "
        "is at risk; any incident involving payment processing.", body))

    flow.append(Paragraph("Communications policy", h2))
    flow.append(Paragraph(
        "All P1/P2 incidents require updates in the #incident channel every "
        "15 minutes. Updates must include: current status, actions taken, "
        "ETA to next action, customer impact estimate. Updates may be brief "
        "but must not be skipped.", body))

    flow.append(Paragraph("Anti-patterns", h2))
    for bullet in [
        "Paging the entire team instead of escalating up the chain.",
        "Silently waiting for the primary to ack past the 8-minute window.",
        "Updating the status page without telling the IC.",
        "Resolving the incident in monitoring before customer impact is confirmed mitigated.",
    ]:
        flow.append(Paragraph("• " + bullet, body))

    doc.build(flow)


# ─────────────────────────────────────────────────────────────────────────────
# 10) on_call_handoff_checklist.pdf
# ─────────────────────────────────────────────────────────────────────────────
def write_handoff_pdf() -> None:
    out = OUT_DIR / "on_call_handoff_checklist.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=LETTER,
                            leftMargin=0.75*inch, rightMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=18, spaceAfter=12)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13, spaceAfter=8, spaceBefore=12)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=11, leading=15, spaceAfter=6)

    flow = [
        Paragraph("On-Call Handoff Checklist", h1),
        Paragraph(
            "Run through this checklist at the start of every shift and again at "
            "the end. The goal is zero-surprise handoffs: the incoming engineer "
            "should not be paged for something the outgoing engineer already knew about.", body),

        Paragraph("At the start of your shift", h2),
    ]
    for item in [
        "Read the last 24 hours of #incident and #ops channels.",
        "Confirm all active alerts are either resolved, acknowledged, or have a tracking ticket.",
        "Check the deploy calendar for upcoming changes during your window.",
        "Verify your paging device is online by sending a test page to yourself.",
        "Confirm VPN, kubectl, AWS console, and Grafana access all work BEFORE you need them.",
        "Review any open postmortem action items that touch services you cover.",
    ]:
        flow.append(Paragraph("☐ " + item, body))

    flow.append(Paragraph("Mid-shift hygiene", h2))
    for item in [
        "After every page: write a one-line note for the next engineer.",
        "Update any runbook you used if a step was wrong or missing.",
        "Close out any silences you opened before they expire silently.",
    ]:
        flow.append(Paragraph("☐ " + item, body))

    flow.append(Paragraph("At the end of your shift", h2))
    for item in [
        "Post a handoff summary in #oncall-handoff including: open incidents, suspicious metrics, planned deploys, anything you 'wouldn't be surprised to see page tonight'.",
        "Acknowledge that the next engineer has seen and accepted the handoff. Do not leave the channel before that.",
        "If you opened any silences during your shift, list them with their expiry times.",
    ]:
        flow.append(Paragraph("☐ " + item, body))

    flow.append(Paragraph("Refusal to hand off", h2))
    flow.append(Paragraph(
        "If an incident is active at the end of your shift, you do NOT hand off "
        "until you have explicit acknowledgement from the incoming engineer that "
        "they are now the incident commander. Silent handoffs during a live "
        "incident are a known cause of customer impact.", body))

    doc.build(flow)


if __name__ == "__main__":
    print(f"Writing corpus to: {OUT_DIR}")
    write_auth_runbook_md()
    write_db_runbook_md()
    write_monitoring_md()
    write_apigw_txt()
    write_payment_txt()
    write_incident_csv()
    write_deployment_docx()
    write_postmortem_docx()
    write_escalation_pdf()
    write_handoff_pdf()

    files = sorted(OUT_DIR.iterdir())
    print(f"\nGenerated {len(files)} documents:")
    for f in files:
        print(f"  {f.name:42s}  {f.stat().st_size:>7,} bytes")
