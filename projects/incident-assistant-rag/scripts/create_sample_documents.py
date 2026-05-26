from pathlib import Path
import csv

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

PROJECT_ROOT = Path(__file__).resolve().parents[1]
SAMPLE_DIR = PROJECT_ROOT / "data" / "sample_documents"


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content.strip() + "\n", encoding="utf-8")


def create_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Deployment Failure SOP", 0)
    sections = {
        "Overview": "This SOP explains how to investigate production issues that begin after a release or deployment.",
        "Post-Deployment Symptoms": "Users cannot log in, API error rate increases, payment timeout appears, health endpoint fails, or latency becomes degraded immediately after deployment.",
        "First Checks": "Check deployment logs, service version, health endpoint, error rate, environment variables, and feature flags.",
        "Rollback Decision": "Prepare rollback if production users are affected or the root cause is not understood quickly.",
        "Escalation": "Escalate to the service owner, backend team, and incident commander when many users are affected.",
    }
    for title, body in sections.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)
    doc.save(path)


def create_pdf(path: Path) -> None:
    styles = getSampleStyleSheet()
    story = []
    sections = [
        ("Database Locks Runbook", "Investigate database locks, blocked queries, active sessions, blocking processes, long-running transactions, and migration-related lock issues."),
        ("Common Symptoms", "Application requests are slow, API calls time out, database CPU rises, and monitoring reports database.production.locks.high."),
        ("First Checks", "Check active sessions, blocked queries, the blocking process, recent migrations, and database connection pool usage."),
        ("Escalation", "Escalate to the DBA team if blocking queries affect production users or if the blocking process is unknown."),
    ]
    for title, body in sections:
        story.append(Paragraph(title, styles["Heading1"]))
        story.append(Paragraph(body, styles["BodyText"]))
        story.append(Spacer(1, 12))
    SimpleDocTemplate(str(path), pagesize=letter).build(story)


def main() -> None:
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    write_text(SAMPLE_DIR / "auth_service_runbook.md", "# Auth Service Runbook\n\nThe auth-service handles user login, JWT token validation, refresh tokens, and session creation. First checks include auth-service health endpoint, auth-service logs, recent deployment, JWT_SECRET, AUTH_DB_URL, and database connectivity. If many production users cannot log in, classify severity as High and escalate to the backend team.")
    write_text(SAMPLE_DIR / "payment_service_latency.txt", "Payment Service Latency Runbook\n\nCheck payment-service logs, external provider status, queue size, retry backlog, database latency, and duplicate charge protection. Payment failures in production are Critical when users are charged or payment risk exists.")
    write_text(SAMPLE_DIR / "monitoring_alerts.md", "# Monitoring Alerts Reference\n\nAlerts include auth.production.login_failure.spike, payment.production.latency.high, database.production.locks.high, and gateway.production.5xx.high. Severity maps to Critical, High, Medium, and Low based on user impact.")
    write_text(SAMPLE_DIR / "api_gateway_5xx_runbook.txt", "API Gateway 5xx Runbook\n\nCheck gateway health, gateway logs, affected route, upstream service health, routing configuration, rate limits, and recent deployment. Escalate High if many production users are affected.")
    write_text(SAMPLE_DIR / "escalation_policy.md", "# Escalation Policy\n\nCritical means all users, production down, payment risk, security, or data loss. High means many users, production login failure, major 5xx spike, or unavailable major service. Escalate to service owner, backend team, DBA, or incident commander based on ownership.")
    with (SAMPLE_DIR / "incident_examples.csv").open("w", encoding="utf-8", newline="") as file:
        writer = csv.writer(file)
        writer.writerow(["incident_id", "service", "environment", "symptom", "severity", "recommended_action", "escalation_team"])
        writer.writerow(["INC-1001", "auth-service", "production", "Many users cannot log in after deployment", "High", "Check auth logs health endpoint deployment logs and JWT environment variables", "backend team"])
        writer.writerow(["INC-1002", "payment-service", "production", "Payment timeout and pending status", "Critical", "Check provider status queue size retry backlog and duplicate charge protection", "payments team"])
    create_docx(SAMPLE_DIR / "deployment_failure_sop.docx")
    create_pdf(SAMPLE_DIR / "database_locks_runbook.pdf")
    print(f"Created sample documents in {SAMPLE_DIR}")


if __name__ == "__main__":
    main()
