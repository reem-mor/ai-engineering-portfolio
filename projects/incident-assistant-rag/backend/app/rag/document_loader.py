import csv
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.core.exceptions import EmptyDocumentError, UnsupportedFileTypeError


class DocumentLoader:
    def load(self, file_path: Path) -> str:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()

        if suffix in {".md", ".txt"}:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".csv":
            text = self._load_csv(file_path)
        elif suffix == ".pdf":
            text = self._load_pdf(file_path)
        elif suffix == ".docx":
            text = self._load_docx(file_path)
        else:
            raise UnsupportedFileTypeError(f"Unsupported file type: {suffix}")

        if not text or not text.strip():
            raise EmptyDocumentError(f"No extractable text found in: {file_path.name}")

        return text

    def _load_csv(self, file_path: Path) -> str:
        rows: list[str] = []
        with file_path.open("r", encoding="utf-8", errors="ignore", newline="") as csv_file:
            reader = csv.DictReader(csv_file)
            if not reader.fieldnames:
                raise EmptyDocumentError(f"CSV has no headers: {file_path.name}")
            for row in reader:
                rows.append(" | ".join(f"{key}: {value}" for key, value in row.items()))
        return "\n".join(rows)

    def _load_pdf(self, file_path: Path) -> str:
        reader = PdfReader(str(file_path))
        parts: list[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts)

    def _load_docx(self, file_path: Path) -> str:
        document = Document(str(file_path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)
